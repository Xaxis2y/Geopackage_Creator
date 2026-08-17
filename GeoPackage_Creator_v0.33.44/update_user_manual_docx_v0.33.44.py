"""Minimal version and FileGDB guidance update for docs/USER_MANUAL.docx."""
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
MANUAL = ROOT / "docs" / "USER_MANUAL.docx"

REPLACEMENTS = {
    2: "Version 0.33.44  |  August 17, 2026",
    16: "The tool requires Python 3.8 or newer and GDAL 3.6 or newer with Python bindings. ISO 19115 metadata validation runs in an isolated helper process so GDAL and lxml/libxml2 do not share native process state. Reportlab is used for PDF reports, and tkinter is included with most Python installations for the GUI.",
    47: "Source FileGDB - click Browse to select the source. Select the File Geodatabase folder itself (do not double-click into it). A .gdb suffix is conventional but not required when the folder contains GEODATABASE_FILE_* markers. Shapefiles and GeoJSON files can be selected as files.",
    146: "Source FileGDB not recognized - select the geodatabase folder itself, not a file inside it. A folder named without .gdb is valid when it contains GEODATABASE_FILE_* markers; verify with ogrinfo if needed.",
}


def main() -> None:
    document = Document(MANUAL)
    if max(REPLACEMENTS) >= len(document.paragraphs):
        raise RuntimeError("Manual paragraph layout changed; refusing update")
    for index, replacement in REPLACEMENTS.items():
        document.paragraphs[index].text = replacement
    if any(document.paragraphs[index].text != replacement for index, replacement in REPLACEMENTS.items()):
        raise RuntimeError("Expected manual text was not found; refusing a partial update")
    document.save(MANUAL)


if __name__ == "__main__":
    main()
