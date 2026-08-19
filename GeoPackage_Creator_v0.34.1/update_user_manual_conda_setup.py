"""Add the required Anaconda Prompt setup flow to the Word user manual."""
# SPDX-License-Identifier: GPL-2.0-or-later
# Copyright (c) 2026 Eui Soo SON
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
MANUAL = ROOT / "docs" / "USER_MANUAL.docx"
REPLACEMENTS = {
    17: "2.2 Setup in Anaconda Prompt (required on Windows)",
    18: "Open Anaconda Prompt, change to the GeoPackage Creator folder, then create and activate the supported environment:",
    19: "conda env create -f environment.yml",
    20: "conda activate geopackage",
    21: "This installs the tested GDAL build and all required Python packages, including lxml, reportlab, pytest, and ttkbootstrap for the GUI. Do not launch the tool from Anaconda's base environment.",
    22: "If the environment already exists, run: conda env update -n geopackage -f environment.yml --prune, then run: conda activate geopackage. You may instead double-click Anaconda_Start.bat to create and activate the same environment interactively.",
    23: "2.3 Manual installation alternatives",
    24: "Manual GDAL and Python dependency instructions are in GDAL_INSTALLATION.txt and INSTALLATION_GUIDE.md. The Conda environment above is the supported setup for Windows.",
}


def main() -> None:
    document = Document(MANUAL)
    if max(REPLACEMENTS) >= len(document.paragraphs):
        raise RuntimeError("Manual paragraph layout changed; refusing update")
    for index, text in REPLACEMENTS.items():
        document.paragraphs[index].text = text
    document.save(MANUAL)


if __name__ == "__main__":
    main()
