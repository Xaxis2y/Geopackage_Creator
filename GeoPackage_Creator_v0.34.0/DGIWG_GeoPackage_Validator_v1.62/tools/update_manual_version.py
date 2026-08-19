"""Minimal in-place version refresh for the retained v1.59 user manual."""
from pathlib import Path
from docx import Document

HERE = Path(__file__).resolve().parent.parent
SOURCE = HERE / "DGIWG_GeoPackage_Validator_User_Manual_v1.60.docx"
OUTPUT = HERE / "DGIWG_GeoPackage_Validator_User_Manual_v1.62.docx"


def replace_in_paragraph(paragraph):
    if "v1.59" not in paragraph.text and "1.59" not in paragraph.text:
        return
    for run in paragraph.runs:
        run.text = run.text.replace("v1.60", "v1.62").replace("1.60", "1.62")


def visit_paragraphs(container):
    for paragraph in container.paragraphs:
        replace_in_paragraph(paragraph)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                visit_paragraphs(cell)


document = Document(SOURCE)
visit_paragraphs(document)
for section in document.sections:
    visit_paragraphs(section.header)
    visit_paragraphs(section.footer)
document.save(OUTPUT)
print(OUTPUT)
